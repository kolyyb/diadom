from flask import Blueprint, jsonify, request, render_template, flash, redirect, url_for, send_file
from app import db
from app.models import Product, ProductHistory, AVAILABLE_PRODUCTS, PRODUCT_NAME_MAPPING
from datetime import datetime
import io
import csv
from sqlalchemy import func
from docx import Document
from docx.shared import Pt
import os
from pathlib import Path

bp = Blueprint('main', __name__)

def format_datetime(value):
    if value is None:
        return ""
    return value.strftime('%d/%m/%Y %H:%M')

@bp.route('/')
def index():
    # Initialiser les produits s'ils n'existent pas
    Product.initialize_products()
    
    # Récupérer tous les produits
    products = Product.query.all()
    
    # Initialiser le dictionnaire des totaux
    product_totals = {}
    for product_name in AVAILABLE_PRODUCTS:
        product_totals[product_name] = {
            'box': 0,
            'apartment': 0,
            'total': 0
        }
    
    # Calculer les totaux
    for product in products:
        if product.quantity is not None:  # Vérifier que la quantité n'est pas None
            product_totals[product.name][product.location] = product.quantity
            product_totals[product.name]['total'] += product.quantity
    
    return render_template('index.html', 
                         products=products,
                         product_totals=product_totals,
                         available_products=AVAILABLE_PRODUCTS,
                         format_datetime=format_datetime)

@bp.route('/api/products', methods=['POST'])
def add_product():
    try:
        product = Product(
            name=request.form['name'],
            quantity=int(request.form['quantity']),
            location=request.form['location']
        )
        db.session.add(product)
        db.session.flush()  # Pour obtenir l'ID du produit
        product.log_change('create')
        db.session.commit()
        flash('Produit ajouté avec succès', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de l\'ajout: {str(e)}', 'danger')
    return redirect(url_for('main.index'))

@bp.route('/api/products/<int:id>', methods=['PUT'])
def update_product(id):
    try:
        data = request.get_json()
        product = Product.query.get_or_404(id)
        
        old_values = {
            'quantity': product.quantity,
            'location': product.location
        }
        
        # Ne mettre à jour que si les valeurs sont différentes
        if (product.quantity != data['quantity'] or 
            product.location != data['location']):
            
            product.quantity = data['quantity']
            product.location = data['location']
            product.last_modified = datetime.utcnow()
            
            product.log_change('update', old_values)
            db.session.commit()
        
        return jsonify({
            'status': 'success',
            'product': {
                'id': product.id,
                'name': product.name,
                'quantity': product.quantity,
                'location': product.location,
                'last_modified': format_datetime(product.last_modified)
            }
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@bp.route('/api/products/<int:id>', methods=['DELETE'])
def delete_product(id):
    try:
        product = Product.query.get_or_404(id)
        db.session.delete(product)
        db.session.commit()
        return '', 204
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400

@bp.route('/api/reset-inventory', methods=['POST'])
def reset_inventory():
    try:
        products = Product.query.all()
        for product in products:
            old_quantity = product.quantity
            product.quantity = 0
            product.last_modified = datetime.utcnow()
            product.log_change('reset', {'quantity': old_quantity})
        db.session.commit()
        flash('Inventaire remis à zéro avec succès', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la remise à zéro: {str(e)}', 'danger')
    return redirect(url_for('main.index'))

@bp.route('/api/export-csv')
def export_csv():
    output = io.StringIO()
    writer = csv.writer(output)
    
    # En-têtes
    writer.writerow(['Nom', 'Quantité', 'Emplacement', 'Dernière modification'])
    
    # Données
    products = Product.query.all()
    for product in products:
        writer.writerow([
            product.name,
            product.quantity,
            product.location,
            product.last_modified.strftime('%Y-%m-%d %H:%M:%S')
        ])
    
    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8')),
        mimetype='text/csv',
        as_attachment=True,
        download_name=f'inventaire_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
    )

@bp.route('/api/generate-order', methods=['GET'])
def generate_order():
    try:
        # Chemin du modèle de bon de commande
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates', 'BON COMMANDE PATIENT_KP100225.docx')
        
        if not os.path.exists(template_path):
            return jsonify({'error': 'Le modèle de bon de commande est introuvable'}), 404

        # Charger le document
        doc = Document(template_path)
        
        # Récupérer les totaux par produit
        products = Product.query.all()
        product_totals = {}
        
        # Initialiser les totaux
        for product_name in AVAILABLE_PRODUCTS:
            product_totals[product_name] = {
                'box': 0,
                'apartment': 0,
                'total': 0
            }
        
        # Calculer les totaux
        for product in products:
            if product.quantity is not None:
                product_totals[product.name][product.location] = product.quantity
                product_totals[product.name]['total'] += product.quantity

        print("\nTotaux calculés pour chaque produit:")
        for name, totals in product_totals.items():
            print(f"  {name}: box={totals['box']}, apartment={totals['apartment']}, total={totals['total']}")

        # Parcourir les tableaux du document
        for table_index, table in enumerate(doc.tables):
            print(f"\nTraitement Table {table_index + 1}")
            for row_index, row in enumerate(table.rows):
                if len(row.cells) < 4:
                    print(f"  Ligne {row_index + 1}: Ignorée (moins de 4 colonnes)")
                    continue

                # Lire le nom du produit dans la 2e colonne
                product_cell_text = row.cells[1].text.strip().upper()
                print(f"  Ligne {row_index + 1}: Produit trouvé dans la 2e colonne: '{product_cell_text}'")
                
                # Chercher une correspondance exacte dans le mapping
                matched_product = None
                for db_name, word_names in PRODUCT_NAME_MAPPING.items():
                    upper_names = [name.upper() for name in word_names]
                    if product_cell_text in upper_names:
                        matched_product = db_name
                        print(f"    -> Correspond à '{db_name}' dans la base")
                        break
                
                if matched_product:
                    # Mettre à jour la 4e colonne avec le total
                    total_cell = row.cells[3]
                    total = product_totals[matched_product]['total']
                    
                    old_text = total_cell.text.strip()
                    total_cell.text = str(total)
                    print(f"    -> Mise à jour 4e colonne: '{old_text}' -> '{total}'")
                    
                    # Appliquer le style
                    paragraph = total_cell.paragraphs[0]
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    font = run.font
                    font.size = Pt(12)
                    font.name = 'Arial'
                    paragraph.alignment = 1  # CENTER
                else:
                    print(f"    -> Aucune correspondance trouvée dans la base")

        # Sauvegarder le document généré
        output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'generated_files')
        os.makedirs(output_path, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = os.path.join(output_path, f'BON_COMMANDE_{timestamp}.docx')
        doc.save(output_file)

        print(f"\nDocument sauvegardé: {output_file}")

        return send_file(
            output_file,
            as_attachment=True,
            download_name=f'BON_COMMANDE_{timestamp}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"\nErreur: {str(e)}")
        return jsonify({'error': str(e)}), 500

@bp.route('/history')
def view_history():
    history = ProductHistory.query.order_by(ProductHistory.timestamp.desc()).all()
    return render_template('history.html', history=history, format_datetime=format_datetime)
